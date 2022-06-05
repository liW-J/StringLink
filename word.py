from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Pt
from function import option

import random,time
#修改六处：①textQueation.txt;②clear.txt;③标题；④clear.txt;⑤key.txt;⑥output.txt

from time import strftime, localtime
#清理数据！
with open("question/textQuestion3.txt", "r", encoding='utf8') as f1:
    with open("./clear/clear3.csv","a",encoding='utf8') as f2:
        for line in f1:
            str1=line.lstrip("{/'history/': [ ")
            str2=str1.rstrip("\']}")

            f2.write(str2)
            # print(str2)

doc1 = Document()  #生成一个空的docx对象
head=doc1.add_paragraph() # 添加标题
head.add_run('二、阅读').font.bold = True # 字形加粗
head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
head.size = Pt(36)
head1=doc1.add_paragraph()
head1.add_run('第一部分').font.bold = True # 字形加粗
head1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
head1.size = Pt(20)
doc1.add_paragraph('选择正确的词语填空') # 添加标题
# p1=doc1.add_paragraph('edit by:HuaBro \t update time: {}'.format(timestr))   # 添加段落
# p1.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中


with open("./clear/clear3.csv","r",encoding='utf8') as f:
    with open("key/key3.csv", "r", encoding='utf8') as f4:
        with open("newKey/newKey3.csv", "r", encoding='utf8') as f5:
            with open("kind/kind3.csv", "r", encoding='utf8') as f6:
                KEY=f5.readlines()
                # Right=f4.readlines()
                keycount=0
                count=0
                for i,j,l in zip(f,f4,f6):
                    print(i)
                    str1=i.split("', '")
                    str2=j.split(",")
                    str3 = l.split(",")
                    keycount+=len(str2)
                    count+=1
                    print(str1)
                    if len(str1)>10:
                        continue
                    # print(str2)
                    #将题目写入docx
                    for k in str1:
                        if k == str1[-1]:
                            k = k.strip("']}\n")
                        doc1.add_paragraph(k)
                    str=""
                    for i in str3:
                        if i == "空":
                            continue
                        else:
                            str+=i+" "
                    doc1.add_paragraph("题目类型:{}".format(str))
                    print(k)
                    for num in range(len(KEY)):

                        if num<keycount and num>=(keycount-len(str2)):

                            print(KEY[num])
                            NEWKEY=KEY[num].split(',')
                            print(NEWKEY)

                            for k in NEWKEY:
                                if k == NEWKEY[-1]:
                                    k = k.strip('\n')
                                    NEWKEY[-1] = k
                                    print(NEWKEY)
                            if len(NEWKEY) == 5:

                                A,B,C,D,RightKey=option(NEWKEY)

                                doc1.add_paragraph("{}.A.{}B.{}C.{}D.{}\t\t正确答案：{},{}".format(len(str2)-(keycount-num)+1,A, B, C, D,RightKey,NEWKEY[-1]))

                            else:
                                A = NEWKEY[0]
                                doc1.add_paragraph("{}. A.{}\t\t正确答案：{}".format(len(str2)-(keycount-num)+1,A,NEWKEY[0]))
     # for l in str2:
                #     doc1.add_paragraph(l)

                # doc1.add_paragraph(j)
        doc1.styles['Normal'].font.name = '宋体'
        doc1.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') # 字体
        doc1.styles['Normal'].font.name = 'Times New Roman' # 数字字体
        doc1.save('./output/output3.docx')
